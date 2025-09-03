import os
import tempfile
from pathlib import Path
from typing import Dict, Any, List, Optional, Union, Tuple
import json
import logging
from datetime import datetime

from worker.model_manager import ModelManager
from worker.storage_manager import StorageManager
from worker.types import JobStatus, DocumentType, ProcessingMode, Document, Page, Bbox
from worker.pipeline.processors.advanced_ocr import AdvancedOCRProcessor
from worker.pipeline.processors.pdf_processor import PDFProcessor
from worker.pipeline.processors.table_detector import TableDetector

from worker.pipeline.pipeline_builder import PipelineBuilder

logger = logging.getLogger(__name__)

class EnhancedDocumentProcessor:

    def __init__(self, model_manager: ModelManager, storage_manager: StorageManager):
        self.model_manager = model_manager
        self.storage_manager = storage_manager
        self.ocr_processor = None
        self.pdf_processor = None
        self.table_detector = None
        self.document_classifier = None
        self.form_extractor = None
        self.layout_analyzer = None

        # Performance controls
        self.fast_mode: bool = True
        self.max_pages: int | None = None

        # Initialize counters and metrics
        self.processed_documents = 0
        self.processing_times = []

        # Load processors on demand to conserve memory
        self._initialize_core_processors()
        self.ocr_processor = None
        self.pdf_processor = None
        self.table_detector = None
        self.document_classifier = None
        self.form_extractor = None
        self.layout_analyzer = None

        # Performance controls
        self.fast_mode: bool = True  # default to fast for local/dev
        self.max_pages: int | None = None

        # Initialize counters and metrics
        self.processed_documents = 0
        self.successful_jobs = 0
        self.failed_jobs = 0
        self.processing_times = []

        # Load processors on demand to conserve memory
        self._initialize_core_processors()

    def _initialize_core_processors(self):
        # Core OCR processor - always needed
        self.ocr_processor = AdvancedOCRProcessor(
            use_paddle=True,
            use_tesseract=True,
            model_path=self.model_manager.get_model_path("ocr")
        )

    def _initialize_pdf_processor(self):
        if not self.pdf_processor:
            dpi = 150 if self.fast_mode else 300
            ocr_threshold = 0.2 if self.fast_mode else 0.1
            self.pdf_processor = PDFProcessor(
                ocr_processor=self.ocr_processor,
                dpi=dpi,
                ocr_threshold=ocr_threshold
            )
            # apply speed-related attributes
            try:
                setattr(self.pdf_processor, "max_pages", self.max_pages)
                setattr(self.pdf_processor, "fast", self.fast_mode)
            except Exception:
                pass

    def _initialize_table_detector(self):
        if not self.table_detector:
            self.table_detector = TableDetector({
                "models_dir": self.model_manager.get_model_path("table_detection")
            })

    def _initialize_document_classifier(self):
        if not self.document_classifier:
            self.document_classifier = DocumentClassifier(
                model_path=self.model_manager.get_model_path("document_classification"),
                feature_extractor_path=self.model_manager.get_model_path("feature_extraction")
            )

    # Removed legacy initializers for form and layout to streamline core pipeline

    def process_document(self, job_id: str, document_path: str, params: Dict[str, Any]) -> Dict[str, Any]:
        start_time = datetime.now()

        logger.info(f"Starting document processing for job {job_id}")
        logger.info(f"Document path: {document_path}")
        logger.info(f"Parameters: {params}")

        try:
            # Parse processing options
            processing_mode = params.get("mode", ProcessingMode.STANDARD)
            # Normalize processing_mode to enum
            if isinstance(processing_mode, str):
                try:
                    processing_mode = ProcessingMode[processing_mode.upper()]
                except Exception:
                    processing_mode = ProcessingMode.STANDARD
            extract_tables = params.get("extract_tables", False)
            classify_document = params.get("classify_document", False)
            extract_forms = params.get("extract_forms", False)
            analyze_layout = params.get("analyze_layout", False)
            output_format = params.get("output_format", "json")

            # Fast-mode controls
            self.fast_mode = bool(params.get("fast", False))  # prefer higher quality by default
            self.max_pages = params.get("max_pages")

            # Mode-driven overrides
            if processing_mode == ProcessingMode.BASIC:
                # aggressive speed
                self.fast_mode = True
                if self.max_pages is None:
                    self.max_pages = 3
                extract_tables = False
                classify_document = False
                extract_forms = False
                analyze_layout = False
            elif processing_mode == ProcessingMode.ADVANCED:
                # full features; allow user to force fast=False by default
                self.fast_mode = bool(params.get("fast", False))
                extract_tables = params.get("extract_tables", True)
                classify_document = False
                extract_forms = False
                analyze_layout = False

            # Determine document type
            doc_type = self._determine_document_type(document_path)
            logger.info(f"Detected document type: {doc_type.name}")

            # Caching: avoid reprocessing if hash+params match
            try:
                with open(document_path, 'rb') as f:
                    content = f.read()
                file_hash = self.storage_manager.compute_hash(content)
                results["cache_key"] = file_hash
            except Exception:
                file_hash = None

            # Initialize base results
            results = {
                "document_type": doc_type.name.lower(),
                "file_path": document_path,
                "file_name": os.path.basename(document_path),
                "processing_mode": processing_mode.name.lower(),
                "processing_params": params,
                "job_id": job_id,
                "timestamp": datetime.now().isoformat(),
                "pages": []
            }

            # Cache short-circuit: reuse prior results if params match
            try:
                allow_cache = bool(params.get("allow_cache", True))
                param_key = {
                    "mode": results["processing_mode"],
                    "extract_tables": params.get("extract_tables", False),
                    "classify_document": params.get("classify_document", False),
                    "extract_forms": params.get("extract_forms", False),
                    "analyze_layout": params.get("analyze_layout", False),
                    "profile": params.get("profile", ""),
                }
                if file_hash and allow_cache:
                    cache_hit = self.storage_manager.cache_get(file_hash, "results.json")
                    if cache_hit and os.path.exists(cache_hit):
                        import json as _json
                        with open(cache_hit, "r", encoding="utf-8") as cf:
                            cached_results = _json.load(cf)
                        crp = cached_results.get("processing_params", {})
                        if all(crp.get(k) == v for k, v in param_key.items()):
                            logger.info(f"Cache hit for job {job_id} (hash={file_hash[:8]}...), reusing results")
                            result_path = self._save_results(job_id, cached_results, output_format)
                            duration = (datetime.now() - start_time).total_seconds()
                            self.processed_documents += 1
                            self.successful_jobs += 1
                            self.processing_times.append(duration)
                            return {
                                "job_id": job_id,
                                "status": JobStatus.COMPLETED,
                                "message": "Document processed from cache",
                                "result_path": result_path,
                                "summary": self._generate_summary(cached_results),
                                "processing_duration": duration,
                                "cache": True,
                            }
            except Exception as _ce:
                logger.debug(f"Cache check skipped: {_ce}")


            # If a pipeline profile is provided, attempt to run the configured processor chain
            profile = params.get("profile")
            # Prefer profile for PDFs automatically to drive ingestion via pipeline
            if doc_type == DocumentType.PDF and not profile:
                profile = "default"

            logger.info(f"Profile selected: {profile}")
            if profile:
                logger.info(f"Starting pipeline execution for profile: {profile}")
                try:
                    pipeline_builder = PipelineBuilder()
                    # Load python config into builder to honor profiles
                    try:
                        from configs.pipeline_config import pipeline_config as _py_pipeline_config
                        pipeline_builder.pipeline_config = {"pipelines": _py_pipeline_config}
                    except Exception as _e:
                        logger.warning(f"Failed to load python pipeline_config: {_e}")
                    pipeline = pipeline_builder.build_pipeline(pipeline_name=profile)

                    export_override = params.get("export_format")
                    output_dir_override = params.get("output_dir")
                    if export_override or output_dir_override:
                        updated = []
                        for proc_id, inst in pipeline:
                            if hasattr(inst, "__class__") and inst.__class__.__name__ == "Exporter":
                                if export_override:
                                    try:
                                        inst.default_format = str(export_override)
                                    except Exception:
                                        pass
                                if output_dir_override:
                                    try:
                                        inst.output_dir = str(output_dir_override)
                                    except Exception:
                                        pass
                            updated.append((proc_id, inst))
                        pipeline = updated

                    # Apply runtime controls (max_pages/fast) to PDFProcessor instances in the pipeline
                    try:
                        for _proc_id, _inst in pipeline:
                            if getattr(_inst, "__class__", None) and _inst.__class__.__name__ == "PDFProcessor":
                                try:
                                    setattr(_inst, "max_pages", self.max_pages)
                                except Exception:
                                    pass
                                try:
                                    setattr(_inst, "fast", self.fast_mode)
                                except Exception:
                                    pass
                    except Exception as _e:
                        logger.debug(f"Failed to apply max_pages/fast to PDFProcessor: {_e}")

                    # Assemble Document for pipeline processors
                    doc = Document(id=results.get("job_id", "doc_cli"), pages=[])
                    # Provide source path and doc type for ingestion processors
                    doc.metadata["source_path"] = document_path
                    doc.metadata["doc_type"] = doc_type.value

                    # Execute processors that accept (document) and return (document)
                    for proc_id, proc_instance in pipeline:
                        process_fn = getattr(proc_instance, "process", None)
                        if callable(process_fn):
                            try:
                                logger.info(f"Running processor: {proc_id}")
                                doc = proc_instance.process(doc)  # type: ignore
                                logger.info(f"After {proc_id}: {len(doc.pages)} pages")
                            except Exception as e:
                                logger.error(f"Processor {proc_id} failed: {e}")
                                continue

                    # Attach regions/tables/tokens for all pages back into results for UI/export
                    logger.info(f"Pipeline complete: {len(doc.pages)} pages to transfer")
                    if doc.pages:
                        page_dicts: List[Dict[str, Any]] = []
                        for p in doc.pages:
                            pd: Dict[str, Any] = {"page_num": getattr(p, "page_num", len(page_dicts) + 1)}
                            regions = getattr(p, "regions", [])
                            tables = getattr(p, "tables", [])
                            tokens = getattr(p, "tokens", [])
                            if regions:
                                pd["regions"] = [
                                    {
                                        "type": r.type,
                                        "bbox": [r.bbox.x1, r.bbox.y1, r.bbox.x2, r.bbox.y2],
                                        "confidence": r.confidence,
                                        "id": r.id,
                                    }
                                    for r in regions
                                ]
                            if tables:
                                pd["tables"] = [
                                    {
                                        "bbox": [t.bbox.x1, t.bbox.y1, t.bbox.x2, t.bbox.y2],
                                        "id": t.id,
                                    }
                                    for t in tables
                                ]
                            if tokens:
                                pd["tokens"] = [
                                    {
                                        "text": tk.text,
                                        "bbox": [tk.bbox.x1, tk.bbox.y1, tk.bbox.x2, tk.bbox.y2],
                                        "confidence": tk.confidence,
                                        "id": tk.id,
                                    }
                                    for tk in tokens
                                ]
                            page_dicts.append(pd)
                        results["pages"] = page_dicts
                        logger.info(f"Transferred {len(page_dicts)} pages to results")
                except Exception as e:
                    logger.warning(f"Profile pipeline execution failed for '{profile}': {e}")
            else:
                # Legacy direct branch by type
                if doc_type == DocumentType.IMAGE:
                    self._process_image(document_path, params, results)
                elif doc_type == DocumentType.TEXT:
                    self._process_text_document(document_path, params, results)
                elif doc_type == DocumentType.DOCX:
                    self._process_docx_document(document_path, params, results)
                else:
                    raise ValueError(f"Unsupported document type: {doc_type}")

            # Additional steps remain available for legacy/direct paths
            if extract_tables and (doc_type == DocumentType.IMAGE):
                # For PDFs, prefer table extraction via pipeline (TableDetector processor)
                self._extract_tables(document_path, params, results)

            # Generate the summary
            results["summary"] = self._generate_summary(results)

            # Calculate processing duration
            end_time = datetime.now()
            duration = (end_time - start_time).total_seconds()
            results["processing_duration"] = duration

            # Save results based on requested output format
            result_path = self._save_results(job_id, results, output_format)

            # Update counters
            self.processed_documents += 1
            self.successful_jobs += 1
            self.processing_times.append(duration)

            logger.info(f"Document processing completed for job {job_id} in {duration:.2f} seconds")

            return {
                "job_id": job_id,
                "status": JobStatus.COMPLETED,
                "message": "Document processed successfully",
                "result_path": result_path,
                "summary": results["summary"],
                "processing_duration": duration
            }

        except Exception as e:
            self.failed_jobs += 1
            logger.error(f"Processing failed for job {job_id}: {str(e)}", exc_info=True)

            end_time = datetime.now()
            duration = (end_time - start_time).total_seconds()

            return {
                "job_id": job_id,
                "status": JobStatus.FAILED,
                "message": f"Processing failed: {str(e)}",
                "error": str(e),
                "processing_duration": duration
            }

    def _determine_document_type(self, document_path: str) -> DocumentType:
        ext = Path(document_path).suffix.lower()

        if ext in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.tif', '.webp']:
            return DocumentType.IMAGE
        elif ext == '.pdf':
            return DocumentType.PDF
        elif ext in ['.doc', '.docx']:
            return DocumentType.DOCX
        elif ext in ['.txt', '.rtf', '.odt', '.html', '.md']:
            return DocumentType.TEXT
        else:
            return DocumentType.UNKNOWN

    def _process_image(self, image_path: str, params: Dict[str, Any], results: Dict[str, Any]) -> None:
        """
        Process an image document.

        Args:
            image_path: Path to the image
            params: Processing parameters
            results: Results dictionary to update
        """
        # Ensure OCR processor is initialized
        if not self.ocr_processor:
            self._initialize_core_processors()

        # Run OCR (downscale for speed in fast mode)
        ocr_input_path = image_path
        if self.fast_mode:
            try:
                import cv2
                import tempfile
                img = cv2.imread(image_path)
                if img is not None:
                    h, w = img.shape[:2]
                    max_dim = max(h, w)
                    target = 1600
                    if max_dim > target:
                        scale = target / float(max_dim)
                        resized = cv2.resize(img, (int(w*scale), int(h*scale)), interpolation=cv2.INTER_AREA)
                        with tempfile.NamedTemporaryFile(suffix=Path(image_path).suffix, delete=False) as tf:
                            cv2.imwrite(tf.name, resized)
                            ocr_input_path = tf.name
            except Exception:
                ocr_input_path = image_path
        ocr_result = self.ocr_processor.process_image(ocr_input_path)

        if "error" in ocr_result:
            raise Exception(ocr_result["error"])

        # Extract text regions
        text_regions = ocr_result.get("results", [])

        # Combine text from all regions
        full_text = " ".join(region.get("text", "") for region in text_regions)

        # Add page info
        page_info = {
            "page_num": 1,
            "text": full_text,
            "text_regions": text_regions,
            "image_path": image_path,
            "size": ocr_result.get("image_size", {})
        }

        results["pages"].append(page_info)
        results["text"] = full_text  # For backward compatibility
        results["confidence"] = self._calculate_avg_confidence(text_regions)

    def _process_pdf(self, pdf_path: str, params: Dict[str, Any], results: Dict[str, Any]) -> None:
        """
        Process a PDF document.

        Args:
            pdf_path: Path to the PDF
            params: Processing parameters
            results: Results dictionary to update
        """
        # Initialize PDF processor if needed
        if not self.pdf_processor:
            self._initialize_pdf_processor()

        # Process the PDF
        pdf_result = self.pdf_processor.process_pdf(pdf_path)

        if "error" in pdf_result:
            raise Exception(pdf_result["error"])

        # Add metadata (skip in fast mode unless explicitly requested)
        if not self.fast_mode or params.get("extract_metadata", False):
            results["metadata"] = self.pdf_processor.extract_pdf_metadata(pdf_path)

        # Process each page
        full_text = ""
        for page_info in pdf_result.get("pages", []):
            page_text = " ".join(block.get("text", "") for block in page_info.get("blocks", []))
            full_text += page_text + "\n\n"

            # Add page to results
            results["pages"].append({
                "page_num": page_info.get("page_num", 0),
                "text": page_text,
                "blocks": page_info.get("blocks", []),
                "is_scanned": page_info.get("is_scanned", False),
                "size": page_info.get("size", {})
            })

        # Add full text for backward compatibility
        results["text"] = full_text.strip()

        # Calculate average confidence score
        confidences = []
        for page in pdf_result.get("pages", []):
            for block in page.get("blocks", []):
                if "confidence" in block:
                    confidences.append(block["confidence"])

        if confidences:
            results["confidence"] = sum(confidences) / len(confidences)
        else:
            results["confidence"] = 1.0  # Assume high confidence for native text

    def _process_text_document(self, document_path: str, params: Dict[str, Any], results: Dict[str, Any]) -> None:
        """
        Process a text-based document.

        Args:
            document_path: Path to the document
            params: Processing parameters
            results: Results dictionary to update
        """
        try:
            # Read the text file
            with open(document_path, 'r', encoding='utf-8') as f:
                text = f.read()

            # Add page info (treating as single page)
            results["pages"].append({
                "page_num": 1,
                "text": text
            })

            # For backward compatibility
            results["text"] = text
            results["confidence"] = 1.0  # Text file has perfect extraction confidence

        except UnicodeDecodeError:
            # Try with different encodings
            encodings = ['latin-1', 'cp1252', 'iso-8859-1']
            for encoding in encodings:
                try:
                    with open(document_path, 'r', encoding=encoding) as f:
                        text = f.read()

                    # Add page info (treating as single page)
                    results["pages"].append({
                        "page_num": 1,
                        "text": text,
                        "encoding": encoding
                    })

                    # For backward compatibility
                    results["text"] = text
                    results["confidence"] = 1.0
                    return
                except UnicodeDecodeError:
                    continue

            # If all encodings fail
            raise Exception("Failed to decode text file with all attempted encodings")

    def _process_docx_document(self, document_path: str, params: Dict[str, Any], results: Dict[str, Any]) -> None:
        """
        Process a DOCX document.

        Args:
            document_path: Path to the document
            params: Processing parameters
            results: Results dictionary to update
        """
        try:
            import docx

            # Open the document
            doc = docx.Document(document_path)

            # Extract text with paragraph breaks
            paragraphs = [p.text for p in doc.paragraphs]
            full_text = "\n".join(paragraphs)

            # Process tables if present
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables.append(table_data)

            # Build structured regions (best-effort without true coordinates)
            text_regions = []
            for p in paragraphs:
                if p and p.strip():
                    text_regions.append({
                        "type": "paragraph",
                        "text": p.strip(),
                    })

            # Add page info (treating as single page since DOCX doesn't have fixed pages)
            results["pages"].append({
                "page_num": 1,
                "text": full_text,
                "paragraphs": paragraphs,
                "text_regions": text_regions,
                "tables": tables
            })

            # For backward compatibility
            results["text"] = full_text
            results["confidence"] = 1.0  # DOCX has perfect extraction confidence

        except ImportError:
            raise Exception("python-docx package is required for DOCX processing")
        except Exception as e:
            raise Exception(f"Failed to process DOCX document: {str(e)}")

    def _extract_tables(self, document_path: str, params: Dict[str, Any], results: Dict[str, Any]) -> None:
        """
        Extract tables from the document.

        Args:
            document_path: Path to the document
            params: Processing parameters
            results: Results dictionary to update
        """
        # Initialize table detector if needed
        if not self.table_detector:
            self._initialize_table_detector()

        doc_type = self._determine_document_type(document_path)

        if doc_type == DocumentType.PDF:
            # For PDFs, process each page
            if not self.pdf_processor:
                self._initialize_pdf_processor()

            # Extract images from PDF if needed
            with tempfile.TemporaryDirectory() as temp_dir:
                img_results = self.pdf_processor.extract_images_from_pdf(
                    document_path, temp_dir)

                # Process tables for each page
                for page_idx, page in enumerate(results["pages"]):
                    # For PDF pages, we need to convert to image first
                    if "image_path" not in page:
                        # Skip pages without images
                        continue

                    # Detect tables in the page
                    tables = self.table_detector.detect_tables(page["image_path"])

                    # Add tables to page results
                    if "tables" not in page:
                        page["tables"] = []

                    page["tables"].extend(tables)

                    # Add table text to each table
                    for table in tables:
                        table_text = self.table_detector.extract_table_text(
                            page["image_path"], table["bbox"])
                        table["data"] = table_text

        elif doc_type == DocumentType.IMAGE:
            # For images, process directly
            tables = self.table_detector.detect_tables(document_path)

            # Add tables to page results
            if results["pages"]:
                if "tables" not in results["pages"][0]:
                    results["pages"][0]["tables"] = []

                results["pages"][0]["tables"].extend(tables)

                # Add table text to each table
                for table in tables:
                    table_text = self.table_detector.extract_table_text(
                        document_path, table["bbox"])
                    table["data"] = table_text

    def _classify_document(self, document_path: str, results: Dict[str, Any]) -> None:
        """
        Classify the document type.

        Args:
            document_path: Path to the document
            results: Results dictionary to update
        """
        # Initialize document classifier if needed
        if not self.document_classifier:
            self._initialize_document_classifier()

        # Get the text content
        text = results.get("text", "")

        # If no text is extracted yet, use the first page
        if not text and results.get("pages"):
            text = results["pages"][0].get("text", "")

        # Load image for visual features if available
        image = None
        doc_type = self._determine_document_type(document_path)
        if doc_type == DocumentType.IMAGE:
            import cv2
            image = cv2.imread(document_path)

        # Classify document
        classification = self.document_classifier.classify(text, image)

        # Extract metadata from text
        metadata = self.document_classifier.extract_document_metadata(text)

        # Add classification results
        results["classification"] = {
            "document_type": classification["document_type"],
            "confidence": classification["confidence"],
            "scores": classification["scores"]
        }

        # Add extracted metadata
        if "metadata" not in results:
            results["metadata"] = {}

        results["metadata"].update(metadata)

    def _extract_form_fields(self, document_path: str, results: Dict[str, Any]) -> None:
        """
        Extract form fields from the document.

        Args:
            document_path: Path to the document
            results: Results dictionary to update
        """
        # Initialize form extractor if needed
        if not self.form_extractor:
            self._initialize_form_extractor()

        # Extract form fields
        form_result = self.form_extractor.extract_form_fields(document_path)

        if "error" in form_result:
            # Log error but continue processing
            logger.warning(f"Form extraction failed: {form_result['error']}")
            return

        # Add form data to results
        results["form_data"] = form_result.get("form_data", {})

    def _analyze_layout(self, document_path: str, results: Dict[str, Any]) -> None:
        """
        Analyze the layout of the document.

        Args:
            document_path: Path to the document
            results: Results dictionary to update
        """
        # Initialize layout analyzer if needed
        if not self.layout_analyzer:
            self._initialize_layout_analyzer()

        doc_type = self._determine_document_type(document_path)

        if doc_type == DocumentType.PDF:
            # For PDFs, analyze each page
            for page_idx, page in enumerate(results["pages"]):
                # If the page has an image path
                if "image_path" in page:
                    layout = self.layout_analyzer.analyze_layout(page["image_path"])
                    page["layout"] = layout

        elif doc_type == DocumentType.IMAGE:
            # For images, analyze directly
            layout = self.layout_analyzer.analyze_layout(document_path)

            # Add layout to the first page
            if results["pages"]:
                results["pages"][0]["layout"] = layout

    def _calculate_avg_confidence(self, text_regions: List[Dict[str, Any]]) -> float:
        """
        Calculate average confidence from text regions.

        Args:
            text_regions: List of text regions with confidence scores

        Returns:
            Average confidence score
        """
        if not text_regions:
            return 0.0

        confidences = []
        for region in text_regions:
            if "confidence" in region and isinstance(region["confidence"], (int, float)):
                confidences.append(region["confidence"])

        if confidences:
            return sum(confidences) / len(confidences)
        else:
            return 0.0

    def _generate_summary(self, results: Dict[str, Any]) -> Dict[str, Any]:
        """
        Generate a summary of the processing results.

        Args:
            results: Processing results

        Returns:
            Summary dictionary
        """
        # Get full text
        text = results.get("text", "")
        if not text and results.get("pages"):
            # Combine text from all pages - check both text field and tokens
            page_texts = []
            for page in results["pages"]:
                page_text = page.get("text", "")
                if not page_text and page.get("tokens"):
                    # Extract text from tokens
                    page_text = " ".join(token.get("text", "") for token in page["tokens"])
                if page_text:
                    page_texts.append(page_text)
            text = " ".join(page_texts)

        words = text.split()

        # Count pages
        page_count = len(results.get("pages", []))

        # Count tables
        table_count = 0
        for page in results.get("pages", []):
            table_count += len(page.get("tables", []))

        # Get document type from classification if available
        doc_class = "unknown"
        if "classification" in results and results["classification"]:
            doc_class = results["classification"].get("document_type", "unknown")

        summary = {
            "word_count": len(words),
            "char_count": len(text),
            "page_count": page_count,
            "table_count": table_count,
            "document_type": results.get("document_type", "unknown"),
            "document_class": doc_class,
            "confidence": results.get("confidence", 0),
        }

        # Add form field count if available
        if "form_data" in results:
            form_fields = results["form_data"].get("form_fields", {})
            summary["form_field_count"] = len(form_fields)

        return summary

    def _save_results(self, job_id: str, results: Dict[str, Any], output_format: str) -> str:
        """
        Save processing results in the requested format.

        Args:
            job_id: Job identifier
            results: Processing results
            output_format: Output format (json, csv, excel)

        Returns:
            Path to the saved results
        """
        # Create base result path
        base_path = self.storage_manager.get_result_path(job_id)

        # Save JSON format (always save for compatibility)
        json_path = f"{base_path}.json"
        with open(json_path, 'w', encoding='utf-8') as f:
            json.dump(results, f, ensure_ascii=False, indent=2)


            # Write to cache for reuse under same hash+params
            try:
                cache_key = results.get("cache_key")
                if cache_key:
                    # Store canonical JSON for cache hit reuse
                    self.storage_manager.cache_put(cache_key, "results.json", json.dumps(results).encode("utf-8"))
            except Exception as _ce:
                logger.debug(f"Cache store skipped: {_ce}")

        # Save in additional formats if requested
        if output_format == "csv":
            self._export_to_csv(results, f"{base_path}.csv")
            return f"{base_path}.csv"

        elif output_format == "excel":
            self._export_to_excel(results, f"{base_path}.xlsx")
            return f"{base_path}.xlsx"

        elif output_format == "text":
            self._export_to_text(results, f"{base_path}.txt")
            return f"{base_path}.txt"

        # Default to JSON
        return json_path

    def _export_to_csv(self, results: Dict[str, Any], output_path: str) -> None:
        """
        Export results to CSV format.

        Args:
            results: Processing results
            output_path: Path to save CSV file
        """
        import csv

        # Create a flattened structure for CSV
        rows = []

        # Process text by page
        for page in results.get("pages", []):
            page_num = page.get("page_num", 0)

            # Process text regions
            for region in page.get("text_regions", []):
                row = {
                    "page": page_num,
                    "text": region.get("text", ""),
                    "confidence": region.get("confidence", 0),
                    "x": region.get("box", [[0, 0]])[0][0],
                    "y": region.get("box", [[0, 0]])[0][1],
                    "type": "text"
                }
                rows.append(row)

            # Process tables
            for table_idx, table in enumerate(page.get("tables", [])):
                for row_idx, row_data in enumerate(table.get("data", [])):
                    for col_idx, cell in enumerate(row_data):
                        row = {
                            "page": page_num,
                            "text": cell,
                            "confidence": 1.0,
                            "table": table_idx + 1,
                            "row": row_idx + 1,
                            "column": col_idx + 1,
                            "type": "table_cell"
                        }
                        rows.append(row)

        # Write CSV file
        if rows:
            with open(output_path, 'w', newline='', encoding='utf-8') as f:
                writer = csv.DictWriter(f, fieldnames=rows[0].keys())
                writer.writeheader()
                writer.writerows(rows)
        else:
            # Create empty CSV with headers
            with open(output_path, 'w', newline='', encoding='utf-8') as f:
                writer = csv.writer(f)
                writer.writerow(["page", "text", "confidence", "type"])

    def _export_to_excel(self, results: Dict[str, Any], output_path: str) -> None:
        """
        Export results to Excel format.

        Args:
            results: Processing results
            output_path: Path to save Excel file
        """
        try:
            import pandas as pd
            from openpyxl import Workbook
            from openpyxl.utils.dataframe import dataframe_to_rows

            # Create workbook with sheets
            wb = Workbook()

            # Summary sheet
            ws_summary = wb.active
            ws_summary.title = "Summary"

            # Add summary info
            summary = results.get("summary", {})
            for i, (key, value) in enumerate(summary.items(), start=1):
                ws_summary.cell(row=i, column=1, value=key)
                ws_summary.cell(row=i, column=2, value=value)

            # Text content sheet
            ws_text = wb.create_sheet("Text Content")

            # Create text dataframe
            text_data = []
            for page in results.get("pages", []):
                text_data.append({
                    "Page": page.get("page_num", 0),
                    "Content": page.get("text", "")
                })

            if text_data:
                df_text = pd.DataFrame(text_data)
                for r_idx, row in enumerate(dataframe_to_rows(df_text, index=False, header=True), 1):
                    for c_idx, value in enumerate(row, 1):
                        ws_text.cell(row=r_idx, column=c_idx, value=value)

            # Tables sheet if tables exist
            tables_exist = False
            for page in results.get("pages", []):
                if page.get("tables"):
                    tables_exist = True
                    break

            if tables_exist:
                ws_tables = wb.create_sheet("Tables")
                row_idx = 1

                for page in results.get("pages", []):
                    page_num = page.get("page_num", 0)

                    for table_idx, table in enumerate(page.get("tables", [])):
                        # Add table header
                        ws_tables.cell(row=row_idx, column=1,
                                     value=f"Page {page_num}, Table {table_idx+1}")
                        row_idx += 2  # Leave a gap

                        # Add table data
                        table_data = table.get("data", [])
                        if table_data:
                            for i, row_data in enumerate(table_data):
                                for j, cell in enumerate(row_data):
                                    ws_tables.cell(row=row_idx + i, column=j+1, value=cell)

                            row_idx += len(table_data) + 2  # Move past table and add gap

            # Form data sheet if form data exists
            if "form_data" in results:
                ws_forms = wb.create_sheet("Form Fields")
                row_idx = 1

                # Form fields
                ws_forms.cell(row=row_idx, column=1, value="Field")
                ws_forms.cell(row=row_idx, column=2, value="Value")
                row_idx += 1

                for field, value in results["form_data"].get("form_fields", {}).items():
                    ws_forms.cell(row=row_idx, column=1, value=field)
                    ws_forms.cell(row=row_idx, column=2, value=value)
                    row_idx += 1

            # Save workbook
            wb.save(output_path)

        except ImportError:
            logger.warning("pandas and openpyxl are required for Excel export")
            # Fallback to CSV
            self._export_to_csv(results, output_path.replace(".xlsx", ".csv"))

    def _export_to_text(self, results: Dict[str, Any], output_path: str) -> None:
        """
        Export results to plain text format.

        Args:
            results: Processing results
            output_path: Path to save text file
        """
        with open(output_path, 'w', encoding='utf-8') as f:
            # Add a title
            f.write(f"Document: {results.get('file_name', 'Unknown')}\n")
            f.write("=" * 80 + "\n\n")

            # Add summary
            f.write("SUMMARY\n")
            f.write("-" * 80 + "\n")
            summary = results.get("summary", {})
            for key, value in summary.items():
                f.write(f"{key}: {value}\n")
            f.write("\n")

            # Add text content by page
            f.write("CONTENT\n")
            f.write("-" * 80 + "\n")

            for page in results.get("pages", []):
                page_num = page.get("page_num", 0)
                f.write(f"\nPAGE {page_num}\n")
                f.write("-" * 40 + "\n")
                f.write(page.get("text", "") + "\n")

                # Add tables
                tables = page.get("tables", [])
                if tables:
                    f.write("\nTABLES\n")
                    for table_idx, table in enumerate(tables):
                        f.write(f"Table {table_idx+1}:\n")

                        # Format table data as text
                        table_data = table.get("data", [])
                        if table_data:
                            # Calculate column widths
                            col_widths = []
                            for row in table_data:
                                while len(col_widths) < len(row):
                                    col_widths.append(0)

                                for i, cell in enumerate(row):
                                    col_widths[i] = max(col_widths[i], len(str(cell)))

                            # Add padding
                            col_widths = [w + 2 for w in col_widths]

                            # Print table
                            for row in table_data:
                                line = ""
                                for i, cell in enumerate(row):
                                    if i < len(col_widths):
                                        line += str(cell).ljust(col_widths[i])
                                f.write(line + "\n")
                        f.write("\n")

            # Add form data if present
            if "form_data" in results:
                f.write("\nFORM FIELDS\n")
                f.write("-" * 80 + "\n")

                for field, value in results["form_data"].get("form_fields", {}).items():
                    f.write(f"{field}: {value}\n")

    def get_processor_stats(self) -> Dict[str, Any]:
        """
        Get statistics about document processor usage.

        Returns:
            Dictionary with processor statistics
        """
        stats = {
            "processed_documents": self.processed_documents,
            "successful_jobs": self.successful_jobs,
            "failed_jobs": self.failed_jobs,
        }

        # Calculate average processing time
        if self.processing_times:
            stats["avg_processing_time"] = sum(self.processing_times) / len(self.processing_times)
            stats["min_processing_time"] = min(self.processing_times)
            stats["max_processing_time"] = max(self.processing_times)
        else:
            stats["avg_processing_time"] = 0
            stats["min_processing_time"] = 0
            stats["max_processing_time"] = 0

        return stats
